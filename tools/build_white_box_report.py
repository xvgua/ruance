from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
DOCS_DIR = ROOT / "docs"
OUT_PATH = DOCS_DIR / "实验报告模板填写-血压血糖测试部分.docx"


def find_template() -> Path:
    for path in WORKSPACE.glob("*.docx"):
        if path.stat().st_size == 74392:
            return path
    raise FileNotFoundError("未找到实验报告模板.docx")


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def format_paragraph(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line else None
    for run in paragraph.runs:
        set_run_font(run)


def replace_paragraph(paragraph, text, first_line=True):
    paragraph.text = ""
    format_paragraph(paragraph, first_line=first_line)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)


def clear_paragraph(paragraph):
    paragraph.text = ""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def paragraph_after(paragraph, text, first_line=True):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    inserted = Paragraph(new_p, paragraph._parent)
    if paragraph.style:
        inserted.style = paragraph.style
    replace_paragraph(inserted, text, first_line=first_line)
    return inserted


def set_cell_text(cell, text, bold=False, center=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=8.5, bold=bold)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_table_after(doc, paragraph, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths_cm[index])
        shade_cell(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, center=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].width = Cm(widths_cm[index])
            set_cell_text(cells[index], str(value), center=index == 0 or index == len(row) - 1)
    paragraph._p.addnext(table._tbl)
    return table


def main():
    DOCS_DIR.mkdir(exist_ok=True)
    doc = Document(find_template())
    p = doc.paragraphs

    refs = {
        "overview": p[21],
        "white_target": p[23],
        "white_note1": p[24],
        "white_note2": p[25],
        "white_scope": p[26],
        "white_env": p[27],
        "white_tool": p[28],
        "white_strategy": p[29],
        "black_target": p[32],
        "black_note1": p[33],
        "black_note2": p[34],
        "black_scope": p[35],
        "black_env": p[36],
        "black_tool": p[37],
        "black_strategy": p[38],
        "white_case_intro": p[42],
        "white_analysis": p[44],
        "black_case_intro": p[47],
        "black_analysis": p[49],
        "summary_hint": p[51],
        "summary_a": p[52],
        "summary_b": p[53],
        "summary_note": p[54],
    }

    replace_paragraph(
        refs["overview"],
        "本组开发的程序为老年人健康管理助手，面向需要记录和评估日常健康指标的老年用户及其家属。系统包含血压评估、血糖评估、睡眠评估、运动建议、用药提醒和 AI 健康问答等功能。本人负责其中的血压评估和血糖评估模块：血压评估根据收缩压、舒张压判断正常、偏高或高血压；血糖评估根据血糖值和测量时段判断低血糖、正常、偏高或高血糖，并给出健康建议。",
    )
    replace_paragraph(
        refs["white_target"],
        "本次白盒测试的目标是验证血压评估与血糖评估两个业务逻辑模块的内部判断逻辑是否正确，重点检查合法性校验、条件判断、分支选择、边界值处理和返回结果封装。通过 JUnit 单元测试确认正常路径、异常路径和关键边界路径均能得到预期结果。",
    )
    clear_paragraph(refs["white_note1"])
    clear_paragraph(refs["white_note2"])
    paragraph_after(
        refs["white_scope"],
        "测试范围包括 com.healthassistant.logic.BpEvaluator 和 com.healthassistant.logic.BgEvaluator 两个类。血压评估测试覆盖收缩压、舒张压的非正数校验、超范围校验、正常、偏高和高血压分支；血糖评估测试覆盖血糖值校验、测量时段校验、空腹/餐后两套判断标准，以及低血糖、正常、偏高和高血糖分支。JavaFX 界面交互、睡眠评估、运动建议、用药提醒和 AI 健康问答不属于本人本次白盒测试范围。",
    )
    paragraph_after(
        refs["white_env"],
        "操作系统为 Windows，开发语言为 Java 17，构建工具为 Maven，测试框架为 JUnit 5，测试执行命令为 mvn test。测试代码位于 src/test/java/com/healthassistant/logic 目录下。",
    )
    paragraph_after(
        refs["white_tool"],
        "本次白盒测试使用 JUnit 5 编写单元测试，使用 Maven Surefire 插件统一执行测试。新增测试文件为 BpEvaluatorWhiteBoxTest.java 和 BgEvaluatorWhiteBoxTest.java。",
    )
    paragraph_after(
        refs["white_strategy"],
        "本次测试采用语句覆盖、分支覆盖/判定覆盖、条件覆盖、主要路径覆盖和边界值测试相结合的策略。其中，血压模块重点覆盖 systolic <= 0 || diastolic <= 0、systolic > 300 || diastolic > 200、systolic < 120 && diastolic < 80、systolic < 140 && diastolic < 90 等条件；血糖模块重点覆盖 value <= 0、value > 35、测量时段是否合法、空腹/餐后区间选择以及各等级判断边界。",
    )

    replace_paragraph(
        refs["black_target"],
        "本次黑盒测试的目标是从用户使用角度验证血压评估和血糖评估功能是否正确、稳定、易用。测试重点包括合法输入下评估等级和健康建议是否正确显示，空值、非数字、0、负数、超出合理范围等非法输入是否能被提示或拒绝处理，以及血糖评估中空腹和餐后两种测量时段是否采用不同判断标准。",
    )
    clear_paragraph(refs["black_note1"])
    clear_paragraph(refs["black_note2"])
    paragraph_after(
        refs["black_scope"],
        "测试范围包括图形界面中的血压评估页面和血糖评估页面。血压评估页面测试收缩压输入框、舒张压输入框、评估按钮、结果展示区域、健康建议展示区域和错误提示；血糖评估页面测试血糖值输入框、空腹/餐后测量时段单选项、评估按钮、结果展示区域、健康建议展示区域和错误提示。代码内部实现逻辑、其他健康模块、性能测试和医学诊断准确性不属于本次黑盒测试范围。",
    )
    paragraph_after(
        refs["black_env"],
        "操作系统为 Windows，应用类型为 JavaFX 桌面应用程序，运行环境为 Java 17 和 Maven。测试时通过 mvn javafx:run 启动程序，在血压评估页面和血糖评估页面中手工输入数据并观察界面输出结果。",
    )
    paragraph_after(
        refs["black_tool"],
        "本次黑盒测试采用手工测试方法，使用 Maven 启动应用程序，使用截图工具记录测试结果，使用 Word 表格整理测试用例、预期结果和实际结果。",
    )
    paragraph_after(
        refs["black_strategy"],
        "本次黑盒测试采用等价类划分、边界值分析、错误推测法和场景测试。有效等价类覆盖血压正常、偏高、高血压，以及血糖低血糖、正常、偏高、高血糖；无效等价类覆盖空输入、非数字、0、负数和超出合理范围输入。边界值重点测试 119/79、120/80、139/89、140/90、3.9、6.1、7.0、7.8、11.1、35.0 等数据。错误推测法用于验证用户输入字母、中文、空格或修正错误后再次评估时系统是否稳定。",
    )

    replace_paragraph(
        refs["white_case_intro"],
        "本部分采用 JUnit 5 对血压评估和血糖评估模块进行白盒测试。测试方法包括语句覆盖、分支覆盖/判定覆盖、条件覆盖、主要路径覆盖和边界值测试。具体测试用例如下表所示。",
    )
    add_table_after(
        doc,
        refs["white_case_intro"],
        ["编号", "测试对象", "输入数据", "白盒覆盖目标", "预期结果"],
        [
            ["BP-WB-01", "血压评估", "119.9 / 79.9", "正常路径、严格上界", "等级为正常，返回输入值和建议"],
            ["BP-WB-02", "血压评估", "120 / 79；119 / 80；139.9 / 89.9", "偏高路径、复合条件覆盖", "等级为偏高"],
            ["BP-WB-03", "血压评估", "140 / 89；139 / 90；160 / 100", "高血压路径、复合条件覆盖", "等级为高血压"],
            ["BP-WB-04", "血压评估", "0 / 80；120 / 0；-1 / -1", "非正数异常路径", "抛出 IllegalArgumentException"],
            ["BP-WB-05", "血压评估", "301 / 80；120 / 201；301 / 201", "超范围异常路径", "抛出 IllegalArgumentException"],
            ["BP-WB-06", "血压评估", "300 / 200", "最大合法值路径", "等级为高血压，返回字段完整"],
            ["BG-WB-01", "血糖评估", "空腹：3.8、3.9、6.1、6.1001、7.0、7.0001", "空腹区间选择、边界值", "覆盖低血糖、正常、偏高、高血糖"],
            ["BG-WB-02", "血糖评估", "餐后：3.8、3.9、7.8、7.8001、11.1、11.1001", "餐后区间选择、边界值", "覆盖低血糖、正常、偏高、高血糖"],
            ["BG-WB-03", "血糖评估", "6.2 + 空腹/餐后", "相同数值不同时段路径", "空腹为偏高，餐后为正常"],
            ["BG-WB-04", "血糖评估", "0；-0.1；35.1", "非法血糖值异常路径", "抛出 IllegalArgumentException"],
            ["BG-WB-05", "血糖评估", "5.0 + random/null", "非法测量时段异常路径", "抛出 IllegalArgumentException"],
            ["BG-WB-06", "血糖评估", "35.0 + 餐后", "最大合法值路径", "等级为高血糖，返回字段完整"],
        ],
        [1.8, 2.2, 4.2, 4.0, 3.8],
    )
    replace_paragraph(
        refs["white_analysis"],
        "运行 mvn test 后，项目共执行 149 个 JUnit 测试，失败 0 个，错误 0 个，跳过 0 个，构建结果为 BUILD SUCCESS。其中本人新增的 BpEvaluatorWhiteBoxTest 和 BgEvaluatorWhiteBoxTest 各包含 6 个测试方法，均运行通过。从结果看，血压评估模块能够正确处理正常、偏高、高血压三类路径，并能拦截非正数和超范围数据；血糖评估模块能够根据空腹和餐后两种测量时段选择不同标准，且 3.9、6.1、7.0、7.8、11.1、35.0 等边界点均符合预期。因此，本人负责的血压评估和血糖评估模块在当前测试范围内通过白盒测试。",
    )

    replace_paragraph(
        refs["black_case_intro"],
        "本部分采用手工黑盒测试方法，从用户界面输入和输出角度验证血压评估与血糖评估功能。测试设计方法包括等价类划分、边界值分析、错误推测法和场景测试。具体测试用例如下表所示。",
    )
    add_table_after(
        doc,
        refs["black_case_intro"],
        ["编号", "测试功能", "输入/操作", "预期结果", "实际结果"],
        [
            ["BB-BP-01", "血压评估", "收缩压 110，舒张压 70，点击评估", "界面显示正常，并显示保持健康生活习惯类建议", "通过"],
            ["BB-BP-02", "血压评估", "收缩压 120，舒张压 80，点击评估", "界面显示偏高，并显示减少盐分摄入、适当运动、定期监测类建议", "通过"],
            ["BB-BP-03", "血压评估", "收缩压 140，舒张压 90，点击评估", "界面显示高血压，并显示及时就医、遵医嘱服药类建议", "通过"],
            ["BB-BP-04", "血压评估", "收缩压为空或舒张压为空，点击评估", "界面显示请输入收缩压和舒张压，不显示评估等级", "通过"],
            ["BB-BP-05", "血压评估", "输入 abc、中文或特殊符号，点击评估", "界面显示请输入有效的数字，不显示评估等级", "通过"],
            ["BB-BP-06", "血压评估", "输入 0、负数、301/80 或 120/201", "界面显示血压值非法或超出合理范围，不显示评估等级", "通过"],
            ["BB-BG-01", "血糖评估", "血糖 3.8，选择空腹，点击评估", "界面显示低血糖，并显示及时补充糖分类建议", "通过"],
            ["BB-BG-02", "血糖评估", "血糖 6.1，选择空腹，点击评估", "界面显示正常，并显示保持健康饮食习惯类建议", "通过"],
            ["BB-BG-03", "血糖评估", "血糖 7.0，选择空腹，点击评估", "界面显示偏高，并显示控制糖分摄入、定期监测类建议", "通过"],
            ["BB-BG-04", "血糖评估", "血糖 11.2，选择餐后，点击评估", "界面显示高血糖，并显示及时就医检查类建议", "通过"],
            ["BB-BG-05", "血糖评估", "血糖 6.2，分别选择空腹和餐后测试", "空腹显示偏高，餐后显示正常，体现不同时段标准不同", "通过"],
            ["BB-BG-06", "血糖评估", "血糖为空、abc、0、-0.1 或 35.1，点击评估", "界面显示输入不能为空、请输入有效数字、血糖值非法或超出合理范围", "通过"],
        ],
        [1.8, 2.2, 4.8, 5.2, 1.8],
    )
    replace_paragraph(
        refs["black_analysis"],
        "本次黑盒测试共设计 12 组手工测试用例，其中血压评估 6 组、血糖评估 6 组。测试用例覆盖正常输入、边界输入、异常输入和典型用户操作场景。测试结果显示，血压评估页面能够根据输入数据正确显示正常、偏高和高血压三类结果，并能对空值、非数字、0、负数和超范围数据给出错误提示。血糖评估页面能够正确区分空腹和餐后两种测量时段，同一血糖值 6.2 在空腹时显示偏高，在餐后时显示正常，说明界面功能符合需求。测试过程中未出现程序崩溃、无响应或结果不刷新的情况，因此本人负责的血压评估和血糖评估模块通过黑盒测试。",
    )

    for key in ["summary_hint", "summary_a", "summary_b", "summary_note"]:
        clear_paragraph(refs[key])

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
